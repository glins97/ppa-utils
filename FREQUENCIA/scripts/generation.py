from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import time
import re
import datetime
from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Inches
from PIL import Image

def number_format(s):
    if s < 10:
        return '0' + str(s)
    return str(s)

def add_data(table, course, content, date):
    indent = '                         '
    size = len(table.rows)
    table.add_row()
    table.cell(size, 0).text = indent + course
    table.cell(size, 1).text = content
    table.cell(size, 2).text = date
    table.cell(size, 0)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="EEEEEE"/>'.format(nsdecls('w'))))
    table.cell(size, 1)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="EEEEEE"/>'.format(nsdecls('w'))))
    table.cell(size, 2)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="EEEEEE"/>'.format(nsdecls('w'))))

def generate_csv():
    # settings -----------------------
    USER_EMAIL = 'romullooliveira@ppa.com'
    USER_PASSWORD = '123456'
    WAIT_TIMEOUT_VALUE = 60
    PAT1 = r'.+? (.+?) (.+?) (.+?)$'
    driver = webdriver.Firefox()
    # --------------------------------

    # do login -----------------------
    driver.get("http://webppa.ddns.net/login")
    username = driver.find_element_by_name("email")
    username.clear()
    username.send_keys(USER_EMAIL)

    password = driver.find_element_by_name("password")
    password.clear()
    password.send_keys(USER_PASSWORD)
    password.send_keys(Keys.RETURN)
    time.sleep(2)
    # --------------------------------

    # get folha ----------------------
    done = False
    data = []
    driver.get('http://webppa.ddns.net/dashboard-select-type')
    WebDriverWait(driver, WAIT_TIMEOUT_VALUE).until(EC.presence_of_element_located((By.CLASS_NAME, 'form-control')))
    form_control = driver.find_element_by_class_name("form-control")
    driver.find_element_by_xpath("//select[@name='type']/option[text()='Por aluno']").click()
    driver.find_element_by_xpath("//button[@name='sender']").click()

    WebDriverWait(driver, WAIT_TIMEOUT_VALUE).until(EC.presence_of_element_located((By.NAME, 'Date')))
    students = driver.find_elements_by_xpath("//select[@name='Item']/option")

    for index, _ in enumerate(students):
        try:
            students = driver.find_elements_by_xpath("//select[@name='Item']/option")
            student = students[index]
            if not student.text: continue
            student_name = student.text
            driver.find_element_by_xpath("//select[@name='Date']/option[text()='Últimos 07 dias']").click()
            driver.find_element_by_xpath("//select[@name='Item']/option[text()=' {} ']".format(student_name)).click()
            driver.find_element_by_xpath("//button[@name='sender']").click()
            time.sleep(1)
            
            table = driver.find_element_by_class_name('table')
            tbody = table.find_element_by_tag_name('tbody')
            rows = tbody.find_elements_by_tag_name('tr')
            data = []
            document = Document('RSF.docx')
            table = document.tables[0]
            added = False
            for row in rows:
                cells = row.find_elements_by_tag_name('td')
                if len(cells) < 5 or not cells[1].text.strip() or cells[3].text.strip() == 'Tutor': continue
                for index, element in enumerate(driver.find_elements_by_class_name('box-success')[:-1]):
                    driver.save_screenshot('inputs/shot.png')
                    location = element.location
                    size = element.size
                    x = location['x']
                    y = location['y']
                    w = size['width']
                    h = size['height']
                    width = x + w
                    height = y + h
                    im = Image.open('inputs/shot.png')
                    im = im.crop((int(x), int(y), int(width), int(height)))
                    im.save("inputs/shot_{}.png".format(index))

                add_data(
                    table,
                    cells[1].text.split('/')[0],
                    ', '.join(cells[1].text.split('/')[1:]),
                    cells[2].text
                )
                added = True

            if added:
                img_table = document.add_table(rows=1, cols=2)
                for i, image in enumerate(['inputs/shot_0.png', 'inputs/shot_1.png']):
                    paragraph = img_table.cell(0, i).paragraphs[0]
                    run = paragraph.add_run()
                    run.add_picture(image, width=Inches(3.9))
                document.save('inputs/{}.docx'.format(student_name.replace(' ', '_').upper()))
            driver.back()
        except Exception as e:
            print(e)

    time.sleep(2)
    driver.close()
